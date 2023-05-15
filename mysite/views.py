# importer les bibliothèques nécessaires
import base64
import urllib
import matplotlib.pyplot as plt
import numpy as np
from django.http import HttpResponse
from django.shortcuts import render
import requests
from bs4 import BeautifulSoup
import re
from PyPDF4 import PdfFileReader
from docx import Document
import html2text
import docx
import io
import sqlite3




#url = request.POST.get('url')
#url = 'https://fr.wikipedia.org/wiki/Liste_des_pays_par_population'
#url = 'https://fr.wikipedia.org/wiki/Liste_des_villes_par_population'
#url = 'https://fr.wikipedia.org/wiki/Population_mondiale'
#url = https://www.mathsisfun.com/data/correlation.html


# Create your views here.

# Create your view here.
def accueil(request):
    return render(request, "accueil.html")

# Create your view here.
def pdf(request):
    return render(request, "pdf.html")

# Create your view here.
def word(request):
    return render(request, "word.html")

# Create your view here.
def defilement(request):
    return render(request, "defilement.html")

'''
def extract_table_from_pdf(file_path, page_num):
    with open(file_path, 'rb') as f:
        pdf = PdfFileReader(f)
        page = pdf.getPage(page_num)
        text = page.extractText()
        table_data = []
        rows = re.findall('\d+\.\d+.*\n', text)
        for row in rows:
            cells = re.findall('\d+\.\d+|[A-Za-z]+', row)
            table_data.append(cells)
        return table_data


def display_table(request):
    table_data = extract_table_from_pdf("C:/Users/jhink/OneDrive/Bureau/tableau.pdf", 0)
    context = {'table_data': table_data}
    return render(request, 'tableauPDF.html', context)
'''

# vue qui extrait le tableau du fichier Word et le renvoie dans un template HTML
def extract_table(request):
    # Chemin vers le fichier Word contenant le tableau
    filepath = 'C:/Users/jhink/OneDrive/Bureau/tableau.docx'

    # Ouvre le fichier Word avec la bibliothèque python-docx
    document = Document(filepath)

    # Récupère le premier tableau du document
    table = document.tables[0]
    data = []

    for row in table.rows:
        data_row = []  # stockage des données d'une ligne
        for cell in row.cells:
            data_row.append(cell.text)
        data.append(data_row)

    # Convertit le tableau en HTML avec la bibliothèque html2text
    html_table = html2text.html2text(table._element.xml)

    # Affiche le tableau dans une page HTML avec le template 'table.html'
    return render(request, 'tableauWord.html', {'html_table': html_table})

def graphiqueW(request):
    doc = docx.Document('C:/Users/jhink/OneDrive/Bureau/tableau.docx')
    table = doc.tables[0]  # première table dans le document
    data1 = []  # stockage des données du tableau
    data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)

    # se connecter à la base de données SQLite3
    conn = sqlite3.connect('C:/Users/jhink/PycharmProjects/djangoProject2/db.sqlite3')
    cursor = conn.cursor()

    # créer une table pour stocker les données du tableau
    cursor.execute('''CREATE TABLE IF NOT EXISTS table_name
                      (column1 TEXT, column2 TEXT)''')

    # insérer les données du tableau dans la table
    for row in data:
        cursor.execute("INSERT INTO table_name (column1, column2) VALUES (?, ?)",
                       (row['Column 1'], row['Column 2']))

    # enregistrer les modifications et fermer la connexion à la base de données
    conn.commit()
    conn.close()

    for row in table.rows:
        data_row = []  # stockage des données d'une ligne
        for cell in row.cells:
            data_row.append(cell.text)
        data1.append(data_row)
    # conversion des données en listes de valeurs
    # première colonne du tableau
    x = [row[0] for row in data1[1:]]
    # deuxième colonne du tableau
    y = [str(row[1].replace(',', '.')) for row in data1[1:]]

    # création du graphique
    plt.plot(x, y)
    # titre de la première colonne
    plt.xlabel(data1[0][0])
    # titre de la deuxième colonne
    plt.ylabel(data1[0][1])
    plt.title('Titre du graphique')

    # génération du code HTML du graphique
    fig = plt.gcf()
    buf = io.BytesIO()
    fig.savefig(buf, format='png')
    buf.seek(0)
    string = base64.b64encode(buf.read())
    uri = urllib.parse.quote(string)
    return render(request, 'graphiqueW.html', {'data': uri})

def graphiqueP(request):
    return render(request, 'graphiqueP.html')

